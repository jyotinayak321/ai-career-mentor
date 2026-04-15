# ================================================
# file_handler.py — File Processing Helper
# ================================================
# YE FILE KYA KARTI HAI?
# PDF aur DOCX files se text extract karti hai
#
# KAB USE HOTI HAI?
# Resume upload hone pe
# parser.py se pehle ye chalti hai
#
# SUPPORTED FORMATS:
# .pdf  → pdfplumber use karta hai
# .docx → python-docx use karta hai
# .txt  → direct decode karta hai
# ================================================

import io
from typing import Tuple


# ------------------------------------------------
# FUNCTION 1 — extract_text_from_pdf()
# ------------------------------------------------
# PDF bytes se text nikalo
# pdfplumber library use karta hai
#
# pdfplumber kyu?
# → Tables bhi handle karta hai
# → Multi-column layout handle karta hai
# → Better than PyPDF2
# ------------------------------------------------

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    PDF file se text extract karo

    Args:
        file_bytes: PDF file ka binary content

    Returns:
        Extracted text string

    Kaise kaam karta hai:
    1. bytes ko file-like object mein convert karo
    2. pdfplumber se open karo
    3. Har page se text nikalo
    4. Combine karke return karo
    """
    try:
        import pdfplumber

        text = ""

        # io.BytesIO = bytes ko file ki tarah treat karo
        # Disk pe save kiye bina!
        with pdfplumber.open(
            io.BytesIO(file_bytes)
        ) as pdf:

            # Har page pe loop karo
            for page_num, page in enumerate(pdf.pages):

                # Page se text nikalo
                page_text = page.extract_text()

                if page_text:
                    text += page_text + "\n"

                # Debug info
                print(
                    f"Page {page_num + 1}: "
                    f"{len(page_text or '')} characters"
                )

        return text.strip()

    except Exception as e:
        print(f"PDF extraction error: {e}")
        raise ValueError(
            f"PDF process nahi ho saki: {str(e)}"
        )


# ------------------------------------------------
# FUNCTION 2 — extract_text_from_docx()
# ------------------------------------------------
# DOCX file se text nikalo
# python-docx library use karta hai
# ------------------------------------------------

def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    DOCX file se text extract karo

    Microsoft Word documents
    python-docx se process hote hain
    """
    try:
        from docx import Document

        # DOCX file open karo
        doc = Document(io.BytesIO(file_bytes))

        text = ""

        # Paragraphs se text nikalo
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"

        # Tables se bhi text nikalo
        # Resume mein tables hote hain!
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + " "
                text += "\n"

        return text.strip()

    except Exception as e:
        print(f"DOCX extraction error: {e}")
        raise ValueError(
            f"DOCX process nahi ho saki: {str(e)}"
        )


# ------------------------------------------------
# FUNCTION 3 — extract_text()
# ------------------------------------------------
# Main function — file type detect karke
# Appropriate function call karo
# ------------------------------------------------

def extract_text(
    file_bytes: bytes,
    filename: str
) -> str:
    """
    File type ke hisaab se text extract karo

    Args:
        file_bytes: File ka binary content
        filename: File ka naam (extension detect karne ke liye)

    Returns:
        Extracted text

    Supported formats:
    .pdf  → PDF extraction
    .docx → DOCX extraction
    .txt  → Direct decode
    """

    filename_lower = filename.lower().strip()

    # PDF file
    if filename_lower.endswith('.pdf'):
        print(f"Processing PDF: {filename}")
        return extract_text_from_pdf(file_bytes)

    # DOCX file
    elif filename_lower.endswith('.docx'):
        print(f"Processing DOCX: {filename}")
        return extract_text_from_docx(file_bytes)

    # TXT file
    elif filename_lower.endswith('.txt'):
        print(f"Processing TXT: {filename}")
        # UTF-8 decode karo
        # errors='ignore' = Invalid chars ignore karo
        return file_bytes.decode(
            'utf-8',
            errors='ignore'
        )

    # Unsupported format
    else:
        extension = filename.split('.')[-1] if '.' in filename else 'unknown'
        raise ValueError(
            f"Unsupported file format: .{extension}\n"
            f"Supported formats: PDF, DOCX, TXT"
        )


# ------------------------------------------------
# FUNCTION 4 — validate_file()
# ------------------------------------------------
# File upload hone se pehle validate karo
# ------------------------------------------------

def validate_file(
    file_bytes: bytes,
    filename: str,
    max_size_mb: int = 5
) -> Tuple[bool, str]:
    """
    File validate karo upload se pehle

    Args:
        file_bytes: File content
        filename: File name
        max_size_mb: Maximum allowed size in MB

    Returns:
        (is_valid, error_message)
        (True, "") = Valid
        (False, "error") = Invalid
    """

    # File size check karo
    # 1 MB = 1024 * 1024 bytes
    max_bytes = max_size_mb * 1024 * 1024
    file_size_mb = len(file_bytes) / (1024 * 1024)

    if len(file_bytes) > max_bytes:
        return False, (
            f"File size {file_size_mb:.1f}MB — "
            f"Maximum {max_size_mb}MB allowed!"
        )

    # File extension check karo
    allowed_extensions = ['.pdf', '.docx', '.txt']
    filename_lower = filename.lower()

    if not any(
        filename_lower.endswith(ext)
        for ext in allowed_extensions
    ):
        return False, (
            f"File type allowed nahi hai! "
            f"Sirf {', '.join(allowed_extensions)} allowed hain."
        )

    # File empty toh nahi?
    if len(file_bytes) < 100:
        return False, "File bahut choti hai — empty file upload ki?"

    return True, ""


# ------------------------------------------------
# FUNCTION 5 — get_file_info()
# ------------------------------------------------
# File ki basic information lo
# ------------------------------------------------

def get_file_info(
    file_bytes: bytes,
    filename: str
) -> dict:
    """
    File ki information return karo
    """
    size_bytes = len(file_bytes)
    size_kb = size_bytes / 1024
    size_mb = size_kb / 1024

    # Extension nikalo
    extension = filename.split('.')[-1].upper() \
        if '.' in filename else 'UNKNOWN'

    return {
        "filename": filename,
        "extension": extension,
        "size_bytes": size_bytes,
        "size_kb": round(size_kb, 2),
        "size_mb": round(size_mb, 2),
        "size_display": (
            f"{size_mb:.2f} MB"
            if size_mb >= 1
            else f"{size_kb:.1f} KB"
        )
    }